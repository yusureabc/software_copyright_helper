import sys
import os
from docx import  Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches

# Usage: python3 helper.py directory suffix file_name
# Command: python3 helper.py /mnt/hgfs/www/test/java java demo

directory = sys.argv[1]
suffix    = sys.argv[2]
file_name = sys.argv[3]

suffix_set = suffix.split( ',' )

# Open doc
document = Document()
# Add text
paragraph = document.add_paragraph( '' )


# 循环读取 目录下的所有文件，如果是所需要的格式，读取内容追加到 docx，最后保存文件
for fpathe,dirs,fs in os.walk( directory ):
  for f in fs:
    # os.path.splitext
    full_file_path = os.path.join( fpathe, f )
    res = os.path.splitext( full_file_path )
    # 取出后缀
    file_suffix = res[1]
    file_suffix = file_suffix.replace( '.', '' )
    if file_suffix in suffix_set:
        try:
            f = open( full_file_path, 'r' )
            file_content = f.read()
            print( file_content )

            run = paragraph.add_run( file_content + '\n\n' )
            run.font.size = Pt( 12 )
            # Set font
            run.font.name = 'Consolas'
        finally:
            if f:
                f.close()

# Save file
document.save( file_name + '.docx')